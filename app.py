import os
import io
import json
import random
import secrets
import smtplib
import base64
import mimetypes
import threading
from datetime import datetime, date, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.utils import formataddr
from functools import wraps
from flask import Flask, render_template, request, jsonify, redirect, session, abort
from groq import Groq
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from tinydb import TinyDB, Query
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

# Allow HTTP for local dev; Vercel uses HTTPS automatically
if not os.getenv('VERCEL'):
    os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', os.urandom(24).hex())
app.permanent_session_lifetime = timedelta(days=30)

# On Vercel, filesystem is read-only except /tmp
DB_DIR  = '/tmp/database' if os.getenv('VERCEL') else 'database'
DB_PATH = os.path.join(DB_DIR, 'db.json')
os.makedirs(DB_DIR, exist_ok=True)
db             = TinyDB(DB_PATH)
users_table    = db.table('users')
settings_table = db.table('settings')
records_table  = db.table('records')
tickets_table  = db.table('tickets')
otps_table     = db.table('otps')

ADMIN_EMAIL       = 'wp113.department@gmail.com'
DRIVE_SCOPES      = ['https://www.googleapis.com/auth/drive.file',
                     'https://www.googleapis.com/auth/gmail.send',
                     'https://www.googleapis.com/auth/gmail.settings.basic']
DRIVE_FOLDER_NAME = 'Daily Work Updates'
DB_SYNC_FOLDER    = 'reporting_users'
MAX_ATTACHMENT_BYTES = 3 * 1024 * 1024  # 3 MB raw — base64 (~1.33x) + JSON overhead stays under Vercel's 4.5MB request body cap

def _get_creds_path():
    """Return path to credentials.json — from file or GOOGLE_CREDENTIALS_JSON env var."""
    if os.path.exists('credentials.json'):
        return 'credentials.json'
    creds_env = os.getenv('GOOGLE_CREDENTIALS_JSON', '')
    if creds_env:
        tmp = '/tmp/credentials.json'
        with open(tmp, 'w') as f:
            f.write(creds_env)
        return tmp
    return None

def _get_redirect_uri():
    host = request.host
    if 'localhost' in host or '127.0.0.1' in host:
        return f'http://{host}/oauth2callback'
    return f'https://{host}/oauth2callback'

# ─── OTP STORE ───────────────────────────────────────────────────────────────
otp_store = {}  # in-memory fallback

def generate_otp(email):
    otp    = str(random.randint(100000, 999999))
    expiry = (datetime.now() + timedelta(minutes=5)).isoformat()
    # Store in both memory and DB so it survives across Vercel instances
    otp_store[email] = {'otp': otp, 'expiry': datetime.now() + timedelta(minutes=5)}
    O = Query()
    if otps_table.search(O.email == email):
        otps_table.update({'otp': otp, 'expiry': expiry}, O.email == email)
    else:
        otps_table.insert({'email': email, 'otp': otp, 'expiry': expiry})
    print(f"\n{'='*52}")
    print(f"  OTP for {email}:  {otp}")
    print(f"  Expires in 5 minutes")
    print(f"{'='*52}\n")
    return otp

def verify_otp(email, otp):
    # Check in-memory first, then DB (handles cross-instance on Vercel)
    now = datetime.now()
    if email in otp_store:
        stored = otp_store[email]
        if now <= stored['expiry'] and stored['otp'] == otp:
            otp_store.pop(email, None)
            otps_table.remove(Query().email == email)
            return True
        if now > stored['expiry']:
            otp_store.pop(email, None)
    # Fallback: check DB
    O     = Query()
    found = otps_table.search(O.email == email)
    if not found:
        return False
    stored = found[0]
    try:
        exp = datetime.fromisoformat(stored['expiry'])
    except Exception:
        otps_table.remove(O.email == email)
        return False
    if now > exp:
        otps_table.remove(O.email == email)
        return False
    if stored['otp'] != otp:
        return False
    otps_table.remove(O.email == email)
    otp_store.pop(email, None)
    return True


# ─── DRIVE DB SYNC ───────────────────────────────────────────────────────────
_startup_synced = False

def _bootstrap_admin_token():
    """On a fresh deploy the DB is empty. If ADMIN_DRIVE_TOKEN_JSON env var is set,
    inject it into the DB so sync_db_from_drive() can then restore everything from Drive."""
    token_json = os.getenv('ADMIN_DRIVE_TOKEN_JSON', '').strip()
    if not token_json:
        return
    S = Query()
    found = settings_table.search(S.email == ADMIN_EMAIL)
    if found and found[0].get('drive_token_json'):
        return  # already in DB from previous sync
    try:
        creds = Credentials.from_authorized_user_info(json.loads(token_json), DRIVE_SCOPES)
        # Use token/expiry check directly — creds.valid also checks scope match
        # which fails for tokens that predate our gmail scopes, causing false negatives
        if not creds.token or creds.expired:
            if creds.refresh_token:
                creds.refresh(Request())
                token_json = creds.to_json()
        if found:
            settings_table.update({'drive_token_json': token_json}, S.email == ADMIN_EMAIL)
        else:
            settings_table.insert({'email': ADMIN_EMAIL, 'drive_token_json': token_json})
        print("🔑  Admin token bootstrapped from ADMIN_DRIVE_TOKEN_JSON env var")
    except Exception as e:
        print(f"⚠️  Could not bootstrap admin token: {e}")

def _get_admin_drive_service_safe():
    try:
        return get_user_drive_service(ADMIN_EMAIL)
    except Exception:
        return None

def _reload_db(content_str):
    global db, users_table, settings_table, records_table, tickets_table, otps_table
    db.close()
    with open(DB_PATH, 'w', encoding='utf-8') as f:
        f.write(content_str)
    db             = TinyDB(DB_PATH)
    users_table    = db.table('users')
    settings_table = db.table('settings')
    records_table  = db.table('records')
    tickets_table  = db.table('tickets')
    otps_table     = db.table('otps')

def sync_db_to_drive():
    """Push local db.json → Drive/reporting_users/ after every write."""
    try:
        service = _get_admin_drive_service_safe()
        if not service:
            return
        folder_id = get_or_create_folder(service, DB_SYNC_FOLDER)
        with open(DB_PATH, 'rb') as f:
            upload_to_drive(service, folder_id, 'db.json', f.read(), 'application/json')
        print("☁️  DB synced to Drive")
    except Exception as e:
        print(f"⚠️  DB sync to Drive failed: {e}")

def sync_db_from_drive():
    """Pull Drive/reporting_users/db.json → local on startup / login.
    Calls _bootstrap_admin_token() first so fresh deploys can self-heal."""
    _bootstrap_admin_token()
    try:
        service = _get_admin_drive_service_safe()
        if not service:
            print("⚠️  sync_db_from_drive: no Drive service (bootstrap token issue?)")
            return
        print("☁️  sync_db_from_drive: Drive service OK, searching for DB folder...")
        q = (f"name='{DB_SYNC_FOLDER}' and "
             f"mimeType='application/vnd.google-apps.folder' and trashed=false")
        folders = service.files().list(q=q, fields='files(id)').execute().get('files', [])
        if not folders:
            print(f"⚠️  sync_db_from_drive: folder '{DB_SYNC_FOLDER}' not found in Drive")
            return
        q2 = f"name='db.json' and '{folders[0]['id']}' in parents and trashed=false"
        files = service.files().list(q=q2, fields='files(id)').execute().get('files', [])
        if not files:
            print("⚠️  sync_db_from_drive: db.json not found in Drive folder")
            return
        raw = service.files().get_media(fileId=files[0]['id']).execute()
        text = raw.decode('utf-8') if isinstance(raw, bytes) else raw
        parsed = json.loads(text)        # validate before overwriting
        settings_count = len(parsed.get('settings', {}))
        users_count    = len(parsed.get('users', {}))
        print(f"☁️  DB synced FROM Drive — {users_count} users, {settings_count} settings rows")
        _reload_db(text)
        # Re-inject bootstrap token after reload in case Drive DB had no admin token.
        # Without this, sync_db_to_drive() would fail when the user saves settings.
        _bootstrap_admin_token()
    except Exception as e:
        print(f"⚠️  DB sync from Drive failed: {e}")

@app.before_request
def startup_sync():
    global _startup_synced
    if not _startup_synced:
        _startup_synced = True
        # Synchronous on Vercel (serverless — background threads are unreliable).
        # Background thread on Railway (persistent server — avoids blocking first request).
        if os.getenv('VERCEL'):
            sync_db_from_drive()
        else:
            threading.Thread(target=sync_db_from_drive, daemon=True).start()


# ─── APPROVAL HELPERS ────────────────────────────────────────────────────────
def get_user_status(email):
    U = Query()
    found = users_table.search(U.email == email)
    return found[0].get('status', 'unknown') if found else 'unknown'

def is_user_approved(email):
    if email == ADMIN_EMAIL:
        return True
    return get_user_status(email) == 'approved'

def _make_attachment_part(attachment):
    """Build a MIME part for an email attachment dict: {filename, mime_type, bytes}."""
    filename  = attachment['filename']
    mime_type = attachment.get('mime_type') or mimetypes.guess_type(filename)[0] or 'application/octet-stream'
    maintype, _, subtype = mime_type.partition('/')
    part = MIMEBase(maintype or 'application', subtype or 'octet-stream')
    part.set_payload(attachment['bytes'])
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', 'attachment', filename=filename)
    return part

def _decode_attachment(attachment_in):
    """Decode + size-validate a client {filename, mime_type, data_base64} payload.
    Returns (attachment_dict_or_None, error_message_or_None)."""
    if not attachment_in or not attachment_in.get('data_base64'):
        return None, None
    try:
        file_bytes = base64.b64decode(attachment_in['data_base64'])
    except Exception:
        return None, 'Invalid attachment data.'
    if len(file_bytes) > MAX_ATTACHMENT_BYTES:
        return None, 'Attachment too big! Max size is 3 MB.'
    return {
        'filename':  attachment_in.get('filename') or 'attachment',
        'mime_type': attachment_in.get('mime_type') or '',
        'bytes':     file_bytes,
    }, None

def send_via_gmail_api(sender_email, to_emails, subject, html_body, attachment=None, sender_name=None):
    """Send email via Gmail API using raw urllib — no discovery doc download needed."""
    import urllib.request as _req, urllib.error
    S = Query()
    found = settings_table.search(S.email == sender_email)
    if not found or not found[0].get('drive_token_json'):
        raise Exception("Drive not connected. Go to Settings → Connect Google Drive.")
    creds = Credentials.from_authorized_user_info(
        json.loads(found[0]['drive_token_json']), DRIVE_SCOPES)
    if not creds.token or creds.expired:
        if creds.refresh_token:
            creds.refresh(Request())
            settings_table.update({'drive_token_json': creds.to_json()}, S.email == sender_email)
        else:
            raise Exception("Drive session expired. Please reconnect in Settings.")
    # Try to fetch Gmail signature (requires gmail.settings.basic.readonly scope)
    try:
        sig_req = _req.Request(
            f'https://gmail.googleapis.com/gmail/v1/users/me/settings/sendAs/{sender_email}',
            headers={'Authorization': f'Bearer {creds.token}'},
            method='GET'
        )
        sig_data = json.loads(_req.urlopen(sig_req, timeout=10).read())
        sig_html = sig_data.get('signature', '').strip()
        if sig_html:
            html_body = html_body + f'<br><br><div>--<br>{sig_html}</div>'
    except Exception:
        pass  # no signature or scope not granted yet — send without it

    if attachment:
        msg = MIMEMultipart('mixed')
        alt = MIMEMultipart('alternative')
        alt.attach(MIMEText(html_body, 'html'))
        msg.attach(alt)
    else:
        msg = MIMEMultipart('alternative')
        msg.attach(MIMEText(html_body, 'html'))
    msg['Subject'] = subject
    if sender_name:
        msg['From'] = formataddr((sender_name, sender_email))
    else:
        msg['From'] = sender_email
    msg['To']      = to_emails[0] if to_emails else ''
    if len(to_emails) > 1:
        msg['Cc'] = ', '.join(to_emails[1:])
    if attachment:
        msg.attach(_make_attachment_part(attachment))
    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
    payload = json.dumps({'raw': raw}).encode('utf-8')
    req = _req.Request(
        'https://gmail.googleapis.com/gmail/v1/users/me/messages/send',
        data=payload,
        headers={'Authorization': f'Bearer {creds.token}',
                 'Content-Type': 'application/json'},
        method='POST'
    )
    try:
        _req.urlopen(req, timeout=15)
    except urllib.error.HTTPError as e:
        body = e.read().decode('utf-8', errors='replace')
        raise Exception(f"Gmail API {e.code}: {body}")

def _send_raw_email(from_user, from_pass, to_email, subject, html_body, cc_emails=None, attachment=None, sender_name=None):
    all_rcpt = [to_email] + (cc_emails or [])
    last_err = None

    # Try Gmail API first (works on Railway via HTTPS, appears in Sent folder)
    try:
        send_via_gmail_api(ADMIN_EMAIL, all_rcpt, subject, html_body, attachment=attachment, sender_name=sender_name)
        return  # success
    except Exception as e:
        last_err = e  # Drive not connected or scope missing — try SMTP next

    # Try SMTP port 587 STARTTLS (less restricted than 465 SSL on cloud servers)
    if from_user and from_pass:
        try:
            if attachment:
                msg = MIMEMultipart('mixed')
                alt = MIMEMultipart('alternative')
                alt.attach(MIMEText(html_body, 'html'))
                msg.attach(alt)
            else:
                msg = MIMEMultipart('alternative')
                msg.attach(MIMEText(html_body, 'html'))
            msg['Subject'] = subject
            if sender_name:
                msg['From'] = formataddr((sender_name, from_user))
            else:
                msg['From'] = from_user
            msg['To']      = to_email
            if cc_emails:
                msg['Cc'] = ', '.join(cc_emails)
            if attachment:
                msg.attach(_make_attachment_part(attachment))
            with smtplib.SMTP('smtp.gmail.com', 587, timeout=8) as sv:
                sv.ehlo()
                sv.starttls()
                sv.login(from_user, from_pass)
                sv.sendmail(from_user, all_rcpt, msg.as_string())
            return  # success
        except Exception as e:
            last_err = e

    raise last_err or Exception('No email credentials available')

def send_approval_request(new_user_email, approval_url):
    s         = get_user_settings(ADMIN_EMAIL)
    gmail_user = s.get('gmail_user', '')
    gmail_pass = s.get('gmail_app_password', '')
    html = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;">
  <h2 style="color:#4f46e5;margin-bottom:6px;">🔔 New Access Request</h2>
  <p style="color:#64748b;margin-bottom:20px;">Someone wants access to the <strong>Daily Update Generator</strong>.</p>
  <table style="width:100%;border-collapse:collapse;margin-bottom:24px;">
    <tr>
      <td style="padding:10px 14px;font-weight:bold;background:#f1f5f9;border:1px solid #e2e8f0;width:140px;">User Email</td>
      <td style="padding:10px 14px;border:1px solid #e2e8f0;">{new_user_email}</td>
    </tr>
    <tr>
      <td style="padding:10px 14px;font-weight:bold;background:#f1f5f9;border:1px solid #e2e8f0;">Requested At</td>
      <td style="padding:10px 14px;border:1px solid #e2e8f0;">{datetime.now().strftime('%d/%m/%Y %H:%M')}</td>
    </tr>
  </table>
  <a href="{approval_url}"
     style="display:inline-block;background:#10b981;color:white;padding:14px 36px;border-radius:8px;text-decoration:none;font-weight:bold;font-size:16px;letter-spacing:0.3px;">
    ✅ Approve Access
  </a>
  <p style="color:#94a3b8;font-size:12px;margin-top:20px;">
    Clicking above will approve this user and automatically send them a confirmation email.
  </p>
</div>"""
    if gmail_user and gmail_pass:
        try:
            _send_raw_email(gmail_user, gmail_pass, ADMIN_EMAIL,
                            f'Access Request: {new_user_email}', html)
            print(f"✅ Approval request emailed to admin for {new_user_email}")
            return
        except Exception as e:
            print(f"⚠️  Could not email admin: {e}")
    print(f"\n{'='*60}")
    print(f"  APPROVAL NEEDED for: {new_user_email}")
    print(f"  Approve URL: {approval_url}")
    print(f"{'='*60}\n")

def send_approval_confirmation(user_email):
    s         = get_user_settings(ADMIN_EMAIL)
    gmail_user = s.get('gmail_user', '')
    gmail_pass = s.get('gmail_app_password', '')
    html = """
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;padding:24px;">
  <h2 style="color:#10b981;">🎉 અરજી મંજૂર થઈ ગઈ</h2>
  <p>પ્રિય અરજદાર,</p>
  <p>અભિનંદન! 🎉</p>
  <p>આપ ખરેખર નસીબદાર છો કે માનનીય ભાસ્કર બારોટ સાહેબશ્રીએ આપને આ એપ્લિકેશન વાપરવા લાયક ગણ્યા છે અને આપની અરજી મંજૂર કરી છે.</p>
  <p>હવે આપ આ એપ્લિકેશનનો ઉપયોગ કરી શકો છો.</p>
  <p>સાહેબશ્રીએ આપ પર જે વિશ્વાસ મૂક્યો છે, તેની કદર કરશો અને એપ્લિકેશનનો જવાબદારીપૂર્વક ઉપયોગ કરશો.</p>
  <p>સાહેબશ્રીનો હૃદયપૂર્વક આભાર માનશો. 🙏</p>
  <br>
  <p>શુભેચ્છાઓ,<br><strong>ટીમ</strong></p>
</div>"""
    if gmail_user and gmail_pass:
        try:
            _send_raw_email(gmail_user, gmail_pass, user_email,
                            '🎉 Daily Update App — Access Approved', html)
            print(f"✅ Approval confirmation sent to {user_email}")
            return
        except Exception as e:
            print(f"⚠️  Could not send confirmation to {user_email}: {e}")
    print(f"✅ User {user_email} approved (configure admin Gmail to send confirmation emails)")


# ─── SESSION / AUTH ──────────────────────────────────────────────────────────
def get_current_user():
    return session.get('user_email')

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user_email'):
            return redirect('/login')
        return f(*args, **kwargs)
    return decorated

def login_required_api(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('user_email'):
            return jsonify({'error': 'Not authenticated'}), 401
        return f(*args, **kwargs)
    return decorated


# ─── USER / SETTINGS HELPERS ─────────────────────────────────────────────────
def get_user_settings(email):
    S = Query()
    found = settings_table.search(S.email == email)
    return found[0] if found else {}

def save_user_settings(email, data):
    S = Query()
    data['email'] = email
    if settings_table.search(S.email == email):
        settings_table.update(data, S.email == email)
    else:
        settings_table.insert(data)

def user_setup_done(email):
    s = get_user_settings(email)
    return bool(s.get('setup_done'))

def get_user_groq_client(email):
    s = get_user_settings(email)
    key = s.get('groq_api_key') or os.getenv('GROQ_API_KEY', '')
    return Groq(api_key=key)


def get_user_groq_model(email):
    return 'llama-3.3-70b-versatile'



# ─── DRIVE HELPERS (PER-USER) ─────────────────────────────────────────────────
def get_user_drive_service(email):
    S = Query()
    found = settings_table.search(S.email == email)
    if not found:
        raise Exception("Drive not connected. Go to Settings → Connect Google Drive.")
    token_json = found[0].get('drive_token_json', '')
    if not token_json:
        raise Exception("Drive not connected. Go to Settings → Connect Google Drive.")
    creds = Credentials.from_authorized_user_info(json.loads(token_json), DRIVE_SCOPES)
    # Check token and expiry directly — creds.valid also checks scope match which
    # fails for old tokens that predate the gmail scopes we added, causing false negatives
    if not creds.token or creds.expired:
        if creds.refresh_token:
            creds.refresh(Request())
            settings_table.update({'drive_token_json': creds.to_json()}, S.email == email)
        else:
            raise Exception("Drive session expired. Please reconnect in Settings.")
    return build('drive', 'v3', credentials=creds)

def get_or_create_folder(service, name, parent_id=None):
    q = f"name='{name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        q += f" and '{parent_id}' in parents"
    files = service.files().list(q=q, fields='files(id)').execute().get('files', [])
    if files:
        return files[0]['id']
    meta = {'name': name, 'mimeType': 'application/vnd.google-apps.folder'}
    if parent_id:
        meta['parents'] = [parent_id]
    return service.files().create(body=meta, fields='id').execute()['id']

def upload_to_drive(service, folder_id, filename, content_bytes, mime_type):
    q = f"name='{filename}' and '{folder_id}' in parents and trashed=false"
    existing = service.files().list(q=q, fields='files(id)').execute().get('files', [])
    media = MediaIoBaseUpload(io.BytesIO(content_bytes), mimetype=mime_type, resumable=False)
    if existing:
        file_id = existing[0]['id']
        service.files().update(fileId=file_id, media_body=media).execute()
    else:
        file_id = service.files().create(
            body={'name': filename, 'parents': [folder_id]},
            media_body=media, fields='id'
        ).execute()['id']
    return file_id


def save_to_drive(email, date_key, work_date, form_data, teams_message, attachment=None):
    service    = get_user_drive_service(email)
    root_id    = get_or_create_folder(service, DRIVE_FOLDER_NAME)
    user_folder = get_or_create_folder(service, email, root_id)
    teams_fid  = get_or_create_folder(service, 'Teams', user_folder)
    mail_fid   = get_or_create_folder(service, 'Mail',  user_folder)

    t_id = upload_to_drive(service, teams_fid, f"{date_key}.txt",
                            teams_message.encode('utf-8'), 'text/plain')
    m_id = upload_to_drive(service, mail_fid,  f"{date_key}.docx",
                            _build_docx_bytes(email, work_date, form_data, teams_message),
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    result = {
        'teams_drive_id':  t_id,
        'mail_drive_id':   m_id,
        'teams_drive_url': f"https://drive.google.com/file/d/{t_id}/view",
        'mail_drive_url':  f"https://drive.google.com/file/d/{m_id}/view",
    }
    if attachment:
        attach_fid = get_or_create_folder(service, 'Attachments', user_folder)
        # Prefix with date_key so same-named files from different days don't overwrite each other
        drive_name = f"{date_key}_{attachment['filename']}"
        mime_type  = attachment.get('mime_type') or mimetypes.guess_type(attachment['filename'])[0] or 'application/octet-stream'
        a_id = upload_to_drive(service, attach_fid, drive_name, attachment['bytes'], mime_type)
        result['attachment_drive_id']   = a_id
        result['attachment_drive_url']  = f"https://drive.google.com/file/d/{a_id}/view"
        result['attachment_filename']   = attachment['filename']
    return result

def fetch_teams_from_drive(email, file_id):
    service = get_user_drive_service(email)
    content = service.files().get_media(fileId=file_id).execute()
    return content.decode('utf-8') if isinstance(content, bytes) else str(content)

def rebuild_history_from_drive(email):
    """Scan Drive files and rebuild records_table for a user whose history index was lost."""
    service     = get_user_drive_service(email)
    root_id     = get_or_create_folder(service, DRIVE_FOLDER_NAME)
    user_folder = get_or_create_folder(service, email, root_id)
    teams_fid   = get_or_create_folder(service, 'Teams', user_folder)
    mail_fid    = get_or_create_folder(service, 'Mail',  user_folder)

    # List all .txt files in Teams folder
    teams_files = service.files().list(
        q=f"'{teams_fid}' in parents and name contains '.txt' and trashed=false",
        fields='files(id,name,webViewLink)'
    ).execute().get('files', [])

    # List all .docx files in Mail folder — build a name→id/url map
    mail_files = service.files().list(
        q=f"'{mail_fid}' in parents and name contains '.docx' and trashed=false",
        fields='files(id,name,webViewLink)'
    ).execute().get('files', [])
    mail_map = {f['name'].replace('.docx', ''): f for f in mail_files}

    R = Query()
    rebuilt = 0
    for tf in teams_files:
        date_key = tf['name'].replace('.txt', '')
        # Skip if already indexed
        if records_table.search((R.date == date_key) & (R.user == email)):
            continue
        mf       = mail_map.get(date_key, {})
        # date_key like 2026-07-01 → display date 01/07/2026
        try:
            parts        = date_key.split('-')
            display_date = f"{parts[2]}/{parts[1]}/{parts[0]}"
        except Exception:
            display_date = date_key
        record = {
            'user':            email,
            'date':            date_key,
            'display_date':    display_date,
            'teams_drive_id':  tf['id'],
            'mail_drive_id':   mf.get('id', ''),
            'teams_drive_url': f"https://drive.google.com/file/d/{tf['id']}/view",
            'mail_drive_url':  f"https://drive.google.com/file/d/{mf['id']}/view" if mf.get('id') else '',
            'task_count':      0,
            'tickets':         [],
            'saved_at':        date_key,
        }
        records_table.insert(record)
        rebuilt += 1

    if rebuilt:
        sync_db_to_drive()
    return rebuilt


# ─── DOCX BUILDER ────────────────────────────────────────────────────────────
def _build_docx_bytes(email, work_date, form_data, teams_message):
    s = get_user_settings(email)
    dept    = s.get('department', 'Production Team')
    manager = s.get('reporting_manager', 'Vipinraj Nair Sir')

    doc   = Document()
    title = doc.add_heading('Daily Work Update', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, value in [
        ('Date', work_date),
        ('Department', dept),
        ('Reporting Manager', manager),
        ('Total Attendance Hours', f"{form_data.get('attendance_hours','N/A')} Hours"),
        ('Total Productive Hours', f"{form_data.get('productive_hours','N/A')} Hours"),
    ]:
        p   = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    doc.add_paragraph()
    doc.add_heading('Task Summary', level=1)
    tasks = form_data.get('tasks', [])
    if tasks:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, h in enumerate(['Task / Ticket', 'Description of Work Done', 'Productive Hours', 'Status']):
            hdr[i].text = h
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        for task in tasks:
            row = tbl.add_row().cells
            row[0].text = task.get('ticket', 'N/A')
            row[1].text = task.get('description', 'N/A')
            row[2].text = f"{task.get('hours','N/A')} Hours"
            row[3].text = task.get('status', 'N/A')

    meetings = form_data.get('meetings', [])
    doc.add_paragraph()
    doc.add_heading('Meetings', level=1)
    if meetings and any(m.get('name') for m in meetings):
        mt = doc.add_table(rows=1, cols=3)
        mt.style = 'Table Grid'
        mh = mt.rows[0].cells
        for i, h in enumerate(['Meeting Name', 'Duration', 'Purpose']):
            mh[i].text = h
            for para in mh[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        for m in meetings:
            if m.get('name'):
                r = mt.add_row().cells
                r[0].text = m.get('name', 'NA')
                r[1].text = m.get('duration', 'NA')
                r[2].text = m.get('purpose', 'NA')
    else:
        doc.add_paragraph('NA')

    doc.add_paragraph()
    doc.add_heading('Teams Message / Output Summary', level=1)
    for line in teams_message.split('\n'):
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── SIMPLE EMAIL HTML ────────────────────────────────────────────────────────
def build_email_html(work_date, form_data, llm_result, user_email):
    s         = get_user_settings(user_email)
    dept      = s.get('department', 'Production Team')
    manager   = s.get('reporting_manager', 'Vipinraj Nair Sir')
    attendance  = form_data.get('attendance_hours', 'N/A')
    productive  = form_data.get('productive_hours', 'N/A')
    task_summaries = llm_result.get('task_summaries', [])
    meetings   = llm_result.get('meetings', [{'name': 'NA', 'duration': 'NA', 'purpose': 'NA'}])

    task_rows = ""
    for ts in task_summaries:
        pts = "".join(f"<li>{p}</li>" for p in ts.get('description_points', []))
        task_rows += f"""<tr>
<td style="padding:8px;border:1px solid #999;font-weight:bold;vertical-align:top;">{ts['ticket']}</td>
<td style="padding:8px;border:1px solid #999;vertical-align:top;"><ul style="margin:0;padding-left:16px;">{pts}</ul></td>
<td style="padding:8px;border:1px solid #999;text-align:center;vertical-align:top;">{ts.get('hours','N/A')}</td>
<td style="padding:8px;border:1px solid #999;text-align:center;vertical-align:top;">{ts.get('status','N/A')}</td>
</tr>"""

    meet_rows = ""
    for m in meetings:
        meet_rows += f"""<tr>
<td style="padding:8px;border:1px solid #999;">{m.get('name','NA')}</td>
<td style="padding:8px;border:1px solid #999;">{m.get('duration','NA')}</td>
<td style="padding:8px;border:1px solid #999;">{m.get('purpose','NA')}</td>
</tr>"""

    def _meaningful(lst):
        if not lst: return False
        joined = " ".join(str(x).strip().lower() for x in lst)
        return joined not in ('', '-', 'none', 'na', 'n/a') and bool(joined.strip())

    summaries = ""
    for ts in task_summaries:
        def li(items): return "".join(f"<li>{x}</li>" for x in (items or []))

        links_html = ""
        if _meaningful(ts.get('links', [])):
            links_html += "<p><b>Links:</b></p><ul>" + "".join(f"<li>{l}</li>" for l in ts['links']) + "</ul>"
        if _meaningful(ts.get('screenshots_links', [])):
            links_html += "<p><b>Screenshots:</b></p><ul>" + "".join(f"<li>{l}</li>" for l in ts['screenshots_links']) + "</ul>"

        deliverables_label = ts.get('deliverables_label', 'Commits Completed')
        summaries += f"\n<p><b>{ts['ticket']}</b></p>\n{links_html}"
        summaries += f"<p><b>{deliverables_label}:</b></p><ul>{li(ts.get('commits',['N/A']))}</ul>"
        summaries += f"<p><b>Modules/screens/features worked on:</b></p><ul>{li(ts.get('modules',['N/A']))}</ul>"
        summaries += f"<p><b>QA/testing done:</b></p><ul>{li(ts.get('qa_testing',['N/A']))}</ul>"

        if _meaningful(ts.get('documentation', [])):
            summaries += f"<p><b>Documentation completed:</b></p><ul>{li(ts['documentation'])}</ul>"
        if _meaningful(ts.get('blockers', [])):
            summaries += f"<p><b>Blockers / Dependencies:</b></p><ul>{li(ts['blockers'])}</ul>"
        if _meaningful(ts.get('tomorrow', [])):
            summaries += f"<p><b>Tomorrow's Planned Work:</b></p><ul>{li(ts['tomorrow'])}</ul>"
        if _meaningful(ts.get('questions', [])):
            summaries += f"<p><b>Questions:</b></p><ul>{li(ts['questions'])}</ul>"

        summaries += "<hr>"

    return f"""<div style="font-family:Arial,sans-serif;font-size:14px;color:#000;max-width:800px;">
<p>Date: {work_date}</p>
<p>Department: {dept}</p>
<p>Reporting Manager: {manager}</p>
<p>Total Attendance Hours: {attendance} Hours</p>
<p>Productive Hours: {productive} Hours</p>
<br>
<table border="0" cellpadding="0" cellspacing="0" style="width:100%;border-collapse:collapse;">
<thead><tr style="background:#f0f0f0;">
<th style="padding:8px;border:1px solid #999;text-align:left;">Task / Ticket</th>
<th style="padding:8px;border:1px solid #999;text-align:left;">Description of Work Done</th>
<th style="padding:8px;border:1px solid #999;text-align:center;">Productive Hours Spent</th>
<th style="padding:8px;border:1px solid #999;text-align:center;">Status</th>
</tr></thead>
<tbody>{task_rows}</tbody>
</table>
<br>
<table border="0" cellpadding="0" cellspacing="0" style="width:100%;border-collapse:collapse;">
<thead><tr style="background:#f0f0f0;">
<th style="padding:8px;border:1px solid #999;text-align:left;">Meeting Name</th>
<th style="padding:8px;border:1px solid #999;text-align:left;">Duration</th>
<th style="padding:8px;border:1px solid #999;text-align:left;">Purpose</th>
</tr></thead>
<tbody>{meet_rows}</tbody>
</table>
<br><hr>
<p><b>Output Summary:</b></p>
{summaries}
</div>"""


# ─── LLM (GROQ → GEMINI FALLBACK) ───────────────────────────────────────────
def _parse_json(raw):
    if '```' in raw:
        for part in raw.split('```'):
            s = part.strip()
            if s.startswith('json'):
                raw = s[4:].strip(); break
            elif s.startswith('{'):
                raw = s; break
    return json.loads(raw)

def call_llm(form_data, user_email):
    tasks            = form_data.get('tasks', [])
    attendance_hours = form_data.get('attendance_hours', 'N/A')
    productive_hours = form_data.get('productive_hours', 'N/A')
    work_date        = form_data.get('date', get_working_date())
    meetings         = form_data.get('meetings', [])
    s                = get_user_settings(user_email)
    dept             = s.get('department', 'Production Team')
    manager          = s.get('reporting_manager', 'Vipinraj Nair Sir')

    tasks_text = ""
    for i, task in enumerate(tasks, 1):
        def _has_content(v):
            s = (v or '').strip()
            return bool(s) and s.lstrip('-').strip() != ''
        links_val       = (task.get('links', '') or '').strip()
        screenshots_val = (task.get('screenshots_links', '') or '').strip()
        tomorrow_val    = (task.get('tomorrow_task', '') or '').strip()
        questions_val   = (task.get('questions', '') or '').strip()
        tasks_text += f"""
Task {i}:
- Ticket/Project Name: {task.get('ticket','N/A')}
- What I Did Today: {task.get('description','N/A')}
- Hours Spent: {task.get('hours','N/A')}
- Status: {task.get('status','N/A')}
- Links: {links_val if _has_content(links_val) else '[EMPTY - omit Links section]'}
- Screenshots Links: {screenshots_val if _has_content(screenshots_val) else '[EMPTY - omit Screenshots section]'}
- Tomorrow's Task: {tomorrow_val if _has_content(tomorrow_val) else '[EMPTY - omit Tomorrow section]'}
- Questions: {questions_val if _has_content(questions_val) else '[EMPTY - omit Questions section]'}
"""
    if meetings and any(m.get('name') for m in meetings):
        meetings_text = "\n".join(f"Name: {m.get('name','NA')} | Duration: {m.get('duration','NA')} | Purpose: {m.get('purpose','NA')}" for m in meetings)
    else:
        meetings_text = "No meetings today."

    num_tasks = len(tasks)
    prompt = f"""You are a professional AIML developer assistant. Format the daily work update into a detailed, professional JSON report.

WRITING STYLE — follow every rule below:

ACTION VERBS — always use strong, specific verbs. Replace vague words before writing:
  "checked" → reviewed / evaluated / analyzed / examined
  "looked into" → investigated / researched / explored
  "worked on" → developed / designed / built / implemented / optimized
  "did" → conducted / performed / executed
  "went through" → reviewed / studied / analyzed

DESCRIPTION POINTS — outcome-focused, not task repetition:
  Each input bullet becomes ONE full professional sentence (15-25 words).
  Focus on the OUTCOME or finding, not just the activity.
  BAD: "Researched memory types that can be added to enable agents to use memory."
  GOOD: "Identified optimal shared-memory architecture to unify conversation context uniformly across 60+ agents."
  BAD: "Checked available LLMs and added routing."
  GOOD: "Evaluated available LLMs and implemented intelligent routing logic to direct each request to the optimal model."

DELIVERABLES (commits field):
  Concise past-tense action phrases (4-8 words). Outcome-oriented, no vague verbs.
  For research tasks: document outputs, findings, comparisons (e.g. "Documented memory architecture findings", "Produced framework comparison matrix").
  For code tasks: technical actions (e.g. "Implemented LLM routing logic", "Deployed Redis cache layer").

TASK TYPE — classify each task:
  Set "task_type": "research" when the task is R&D, investigation, documentation, comparison, or analysis.
  Set "task_type": "development" when the task is coding, building, deploying, or testing software.
  Set "deliverables_label": "Deliverables Completed" for research tasks, "Commits Completed" for development tasks.

MODULES — specific component/feature/artifact names only. For research: name the specific artifact or area studied.

QA / VALIDATION:
  For research tasks: generate MEANINGFUL validation activities, not fake test entries.
    Good: "Cross-referenced findings with latest official Langgraph documentation."
    Good: "Evaluated framework maturity based on community adoption and benchmark comparisons."
    Bad: "Tested Langgraph integration", "Validated memory type research" (these are meaningless for R&D).
  For development tasks: specific tests actually performed on the code/system.

TOMORROW — expand each point into a professional planning sentence (15-20 words each).
TEAMS BULLETS — full professional sentences (15-25 words), outcome-focused, strong verbs.
BANNED: "I was able to", "Successfully", "In this task I", "I have completed", "Basically", "Was able to", "checked".
If a field says [EMPTY - omit ... section]: skip that section in teams_messages and set [] in JSON.

Date: {work_date}
Department: {dept}
Reporting Manager: {manager}
Total Attendance Hours: {attendance_hours} Hours
Total Productive Hours: {productive_hours} Hours

Tasks:
{tasks_text}

Meetings:
{meetings_text}

CRITICAL: There are {num_tasks} task(s). Return EXACTLY {num_tasks} entries in task_summaries and teams_messages. Never skip or merge tasks.

Return JSON with EXACTLY these fields (no markdown, no extra text):

{{
  "teams_messages": ["full teams message for task 1", "full teams message for task 2"],
  "email_subject": "ATTENDANCE: {work_date}",
  "task_summaries": [
    {{
      "ticket": "TASK NAME",
      "task_type": "research",
      "deliverables_label": "Deliverables Completed",
      "description_points": [
        "Identified optimal shared-memory architecture to unify conversation context across 60+ agents.",
        "Completed deep R&D on memory types and produced a structured reference document.",
        "Evaluated agentic terminology and Langgraph with PyTorch to determine trending memory solutions.",
        "Concluded R&D phase with actionable recommendations for memory feature integration."
      ],
      "hours": "X Hours",
      "status": "In Progress",
      "links": [],
      "screenshots_links": [],
      "commits": ["Documented memory architecture findings", "Produced memory type comparison matrix", "Evaluated Langgraph PyTorch integration options"],
      "modules": ["Shared Memory Architecture", "Langgraph Memory Integration", "PyTorch Compatibility Analysis"],
      "qa_testing": ["Cross-referenced findings with latest official Langgraph documentation.", "Evaluated memory solutions against scalability requirements for 60+ agents.", "Reviewed PyTorch-Langgraph compatibility against project constraints."],
      "documentation": [],
      "blockers": [],
      "tomorrow": [],
      "questions": []
    }}
  ],
  "meetings": [{{"name":"NA","duration":"NA","purpose":"NA"}}]
}}

teams_messages MUST be {num_tasks} string(s). EACH teams_messages[i] FORMAT — use this structure exactly:
Work Update

Date: {work_date}

Project/Client Name - [TICKET NAME]

    ** Worked on [TICKET NAME]: [Status]

    - [Full professional sentence 15-25 words expanding on point 1]
    - [Full professional sentence expanding on point 2]
    - [Full professional sentence expanding on point 3]
    - [Full professional sentence expanding on point 4]
    - [Full professional sentence expanding on point 5]

Actual Hours:
    - [X] Hours

[Include only if Tomorrow's Task is NOT [EMPTY - omit Tomorrow section]:]
Tomorrow's Task:
    - [Expand to professional planning sentence]
    - [Another expanded planning sentence if multiple points]

[Include only if Questions is NOT [EMPTY - omit Questions section]:]
Questions:
    - [question from input]

[Include only if Links is NOT [EMPTY - omit Links section]:]
Links:
    - [link from input]

RULES:
- teams_messages MUST have {num_tasks} items — one per task, never combine
- task_summaries MUST have {num_tasks} items — one per task
- Each input bullet point → one outcome-focused description_points sentence (do not merge or drop any)
- task_type: "research" or "development" per task
- deliverables_label: "Deliverables Completed" for research, "Commits Completed" for development
- commits / deliverables: 5-7 concise outcome-oriented past-tense phrases per task
- modules: 4-7 specific components/features/artifacts (for research: specific topics or documents produced)
- qa_testing: 4-6 meaningful validation sentences — for research: cross-referencing, evaluation, comparison activities, NOT fake "Tested X" entries
- documentation: [] unless explicitly mentioned
- blockers: [] unless blockers mentioned
- links: [] if [EMPTY], otherwise list actual URLs
- screenshots_links: [] if [EMPTY], otherwise list the screenshot URLs as provided
- tomorrow: [] if [EMPTY], otherwise expand each point into professional sentences
- questions: [] if [EMPTY], otherwise include as-is
- OMIT Tomorrow's Task / Questions / Links from teams_messages if their input was [EMPTY]
- No meetings → [{{"name":"NA","duration":"NA","purpose":"NA"}}]
- Return ONLY valid JSON, no markdown fences"""

    client = get_user_groq_client(user_email)
    resp = client.chat.completions.create(
        model=get_user_groq_model(user_email),
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2, max_tokens=4096
    )
    return _parse_json(resp.choices[0].message.content.strip())


# ─── DATE ─────────────────────────────────────────────────────────────────────
def get_working_date():
    today = date.today()
    if today.weekday() == 5:  today -= timedelta(days=1)
    elif today.weekday() == 6: today -= timedelta(days=2)
    return today.strftime('%d/%m/%Y')


# ─── PAGE ROUTES ─────────────────────────────────────────────────────────────
@app.route('/')
@login_required
def index():
    if not user_setup_done(get_current_user()):
        return redirect('/setup')
    return render_template('index.html', user_email=get_current_user())

@app.route('/login')
def login_page():
    if session.get('user_email'):
        return redirect('/')
    return render_template('login.html')

@app.route('/setup')
@login_required
def setup_page():
    s = get_user_settings(get_current_user())
    return render_template('setup.html', settings=s, user_email=get_current_user(),
                           drive_connected=bool(s.get('drive_token_json')),
                           has_gmail_pass=bool(s.get('gmail_app_password')))

@app.route('/logout')
def logout():
    session.clear()
    return redirect('/login')

@app.route('/pending')
def pending_page():
    email = session.get('pending_email', '')
    if not email:
        return redirect('/login')
    if is_user_approved(email):
        session.pop('pending_email', None)
        session.permanent  = True
        session['user_email'] = email
        return redirect('/' if user_setup_done(email) else '/setup')
    return render_template('pending.html', user_email=email)

@app.route('/approve/<token>')
def approve_user(token):
    # Always sync from Drive first — DB may be empty after Railway restart
    sync_db_from_drive()
    U = Query()
    found = users_table.search(U.approval_token == token)
    if not found:
        return render_template('approve_result.html', success=False,
                               message='Invalid or expired approval link.')
    user  = found[0]
    email = user['email']
    if user.get('status') == 'approved':
        return render_template('approve_result.html', success=True,
                               message=f'{email} is already approved.')
    users_table.update({'status': 'approved'}, U.approval_token == token)
    sync_db_to_drive()
    send_approval_confirmation(email)
    return render_template('approve_result.html', success=True,
                           message=f'Access approved for {email}. Confirmation sent.')


# ─── AUTH ROUTES ─────────────────────────────────────────────────────────────
@app.route('/api/auth/send-otp', methods=['POST'])
def api_send_otp():
    email = (request.json or {}).get('email', '').strip().lower()
    if not email or '@' not in email:
        return jsonify({'success': False, 'error': 'Invalid email'}), 400

    # Admin direct login — skip OTP entirely
    if email == ADMIN_EMAIL:
        sync_db_from_drive()
        U = Query()
        if not users_table.search(U.email == email):
            users_table.insert({'email': email, 'created_at': datetime.now().isoformat(),
                                'status': 'approved'})
            threading.Thread(target=sync_db_to_drive, daemon=True).start()
        session.permanent = True
        session['user_email'] = email
        setup_done = user_setup_done(email)
        return jsonify({'success': True, 'bypass': True,
                        'redirect': '/' if setup_done else '/setup'})

    otp = generate_otp(email)  # generates instantly, printed to logs, stored in DB

    # Register user if first time
    U = Query()
    if not users_table.search(U.email == email):
        users_table.insert({'email': email, 'created_at': datetime.now().isoformat(),
                            'status': 'unknown'})

    otp_html = f"""
<div style="font-family:Arial,sans-serif;max-width:480px;margin:0 auto;padding:24px;">
  <h2 style="color:#4f46e5;margin-bottom:4px;">&#128272; Your Login OTP</h2>
  <p style="color:#64748b;margin-bottom:20px;">Daily Update Generator</p>
  <div style="background:#f1f5f9;border-radius:12px;padding:24px;text-align:center;margin-bottom:20px;">
    <div style="font-size:36px;font-weight:900;letter-spacing:10px;color:#1e293b;">{otp}</div>
  </div>
  <p style="color:#64748b;font-size:13px;">This OTP expires in <strong>5 minutes</strong>. Do not share it with anyone.</p>
</div>"""

    def _do_send():
        admin_s    = get_user_settings(ADMIN_EMAIL)
        gmail_user = admin_s.get('gmail_user', '') or os.getenv('ADMIN_GMAIL_USER', '')
        gmail_pass = admin_s.get('gmail_app_password', '') or os.getenv('ADMIN_GMAIL_PASS', '')
        try:
            _send_raw_email(gmail_user, gmail_pass, email,
                            'Your OTP — Daily Update App', otp_html)
            print(f"✅ OTP email delivered to {email}")
        except Exception as e:
            print(f"⚠️  OTP email failed: {e}")

    # On Vercel serverless, threads are unreliable — send synchronously
    if os.getenv('VERCEL'):
        _do_send()
    else:
        threading.Thread(target=_do_send, daemon=True).start()

    return jsonify({'success': True, 'message': 'OTP sent to your email!'})

@app.route('/api/auth/verify-otp', methods=['POST'])
def api_verify_otp():
    data  = request.json or {}
    email = data.get('email', '').strip().lower()
    otp   = data.get('otp', '').strip()

    if not verify_otp(email, otp):
        return jsonify({'success': False, 'error': 'Invalid or expired OTP'}), 400

    # Approval gate — admin always passes
    if not is_user_approved(email):
        status = get_user_status(email)
        if status != 'pending':
            # First access request — generate token and email admin
            token = secrets.token_urlsafe(32)
            U2 = Query()
            users_table.update({'status': 'pending', 'approval_token': token}, U2.email == email)
            sync_db_to_drive()
            approval_url = f"{request.host_url.rstrip('/')}/approve/{token}"
            send_approval_request(email, approval_url)
        session['pending_email'] = email
        return jsonify({'success': True, 'status': 'pending', 'redirect': '/pending'})

    session.permanent = True
    session['user_email'] = email
    setup_done = user_setup_done(email)
    return jsonify({'success': True, 'setup_done': setup_done,
                    'redirect': '/' if setup_done else '/setup'})


# ─── SETTINGS ROUTES ─────────────────────────────────────────────────────────
@app.route('/api/settings/save', methods=['POST'])
@login_required_api
def api_settings_save():
    email = get_current_user()
    data  = request.json or {}
    data['setup_done'] = True
    save_user_settings(email, data)
    sync_db_to_drive()
    return jsonify({'success': True, 'message': 'Settings saved!'})

@app.route('/api/settings/get')
@login_required_api
def api_settings_get():
    email = get_current_user()
    s = get_user_settings(email)
    safe = {k: v for k, v in s.items() if k not in ('gmail_app_password', 'drive_token_json', 'email')}
    has_drive = bool(s.get('drive_token_json'))
    return jsonify({'success': True, 'settings': safe, 'drive_connected': has_drive})


@app.route('/api/debug/token')
@login_required_api
def api_debug_token():
    email = get_current_user()
    S = Query()
    found = settings_table.search(S.email == email)
    if not found or not found[0].get('drive_token_json'):
        return jsonify({'error': 'No token'})
    token_data = json.loads(found[0]['drive_token_json'])
    creds = Credentials.from_authorized_user_info(token_data, DRIVE_SCOPES)
    return jsonify({
        'stored_scopes': token_data.get('scopes'),
        'creds_scopes':  list(creds.scopes) if creds.scopes else None,
        'token_exists':  bool(token_data.get('token')),
        'creds_valid':   creds.valid,
        'creds_expired': creds.expired,
        'has_refresh':   bool(creds.refresh_token),
    })

# ─── DRIVE AUTH ──────────────────────────────────────────────────────────────
@app.route('/api/drive/status')
@login_required_api
def api_drive_status():
    email = get_current_user()
    try:
        get_user_drive_service(email)
        return jsonify({'connected': True})
    except Exception as e:
        return jsonify({'connected': False, 'error': str(e)})

@app.route('/api/validate/groq', methods=['POST'])
@login_required_api
def api_validate_groq():
    data  = request.json or {}
    key   = data.get('api_key', '').strip()
    model = data.get('model', 'llama-3.3-70b-versatile').strip() or 'llama-3.3-70b-versatile'
    if not key:
        return jsonify({'valid': False, 'error': 'No API key provided'})
    try:
        client = Groq(api_key=key)
        client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": "Reply with just: OK"}],
            max_tokens=5, temperature=0
        )
        return jsonify({'valid': True, 'message': f'Groq key is valid! Model "{model}" works.'})
    except Exception as e:
        return jsonify({'valid': False, 'error': str(e)})


@app.route('/api/validate/gmail', methods=['POST'])
@login_required_api
def api_validate_gmail():
    data       = request.json or {}
    gmail_user = data.get('gmail_user', '').strip()
    gmail_pass = data.get('gmail_app_password', '').strip()
    if not gmail_user or not gmail_pass:
        return jsonify({'valid': False, 'error': 'Email and App Password are required'})
    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465, timeout=10) as server:
            server.login(gmail_user, gmail_pass)
        return jsonify({'valid': True, 'message': 'Gmail connection successful!'})
    except smtplib.SMTPAuthenticationError:
        return jsonify({'valid': False, 'error': 'Authentication failed. Check your Gmail address and App Password.'})
    except OSError as e:
        if e.errno == 101:
            # SMTP port blocked on this server — credentials are saved.
            # Emails will be sent via Gmail API using your Google Drive connection.
            return jsonify({'valid': True, 'message': '✅ Credentials saved! Emails will be sent via your Google Drive connection (SMTP not available on this server).'})
        return jsonify({'valid': False, 'error': str(e)})
    except Exception as e:
        return jsonify({'valid': False, 'error': str(e)})


@app.route('/api/drive/start-auth', methods=['POST'])
@login_required_api
def api_drive_start_auth():
    email = get_current_user()
    creds_path = _get_creds_path()
    if not creds_path:
        return jsonify({'success': False, 'error': 'credentials.json not found. Set GOOGLE_CREDENTIALS_JSON env var.'}), 400
    try:
        redirect_uri = _get_redirect_uri()
        flow = Flow.from_client_secrets_file(creds_path, DRIVE_SCOPES,
                                                          redirect_uri=redirect_uri)
        auth_url, state = flow.authorization_url(access_type='offline', prompt='consent')
        session['oauth_mode']        = 'drive'
        session['drive_auth_email']  = email
        session['drive_auth_state']  = state
        return jsonify({'success': True, 'auth_url': auth_url})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/oauth2callback')
def oauth2callback():
    email = session.get('drive_auth_email')
    if not email:
        return redirect('/setup?error=auth_expired')
    creds_path = _get_creds_path()
    if not creds_path:
        return redirect('/setup?error=no_credentials')
    try:
        redirect_uri = _get_redirect_uri()
        flow = Flow.from_client_secrets_file(
            creds_path, DRIVE_SCOPES, redirect_uri=redirect_uri)
        flow.fetch_token(authorization_response=request.url)
        creds = flow.credentials
        token_json = creds.to_json()
        S = Query()
        # Save new token temporarily so sync_db_from_drive() can authenticate
        if settings_table.search(S.email == email):
            settings_table.update({'drive_token_json': token_json}, S.email == email)
        else:
            settings_table.insert({'email': email, 'drive_token_json': token_json})
        # Pull existing Drive DB first — prevents overwriting old data on fresh deploy
        sync_db_from_drive()
        # Re-apply the new token on top of the restored DB
        if settings_table.search(S.email == email):
            settings_table.update({'drive_token_json': token_json}, S.email == email)
        else:
            settings_table.insert({'email': email, 'drive_token_json': token_json})
        sync_db_to_drive()
        session.pop('drive_auth_email', None)
        session.pop('drive_auth_state', None)
        session.pop('oauth_mode', None)
        return redirect('/setup?drive=connected')
    except Exception as e:
        print(f"Drive auth error: {e}")
        return redirect('/setup?error=drive_failed')


# ─── MAIN APP API ─────────────────────────────────────────────────────────────
@app.route('/api/date')
def api_date():
    return jsonify({'date': get_working_date()})

@app.route('/api/tickets')
@login_required_api
def api_tickets():
    email = get_current_user()
    T = Query()
    history_names = set(t['name'] for t in tickets_table.search(T.user == email))
    s = get_user_settings(email)
    saved = [p.strip() for p in s.get('project_names', '').splitlines() if p.strip()]
    all_names = sorted(set(list(history_names) + saved))
    return jsonify({'tickets': all_names, 'pinned': saved})

@app.route('/api/generate', methods=['POST'])
@login_required_api
def api_generate():
    try:
        email     = get_current_user()
        form_data = request.json
        att = (form_data.get('attendance_hours') or '').strip()
        prod = (form_data.get('productive_hours') or '').strip()
        if not att or att == 'N/A':
            return jsonify({'success': False, 'error': 'Total Attendance Hours is required.'}), 400
        if not prod or prod == 'N/A':
            return jsonify({'success': False, 'error': 'Total Productive Hours is required.'}), 400
        if not form_data.get('tasks'):
            return jsonify({'success': False, 'error': 'No tasks provided'}), 400
        for t in form_data['tasks']:
            if not (t.get('description') or '').strip():
                return jsonify({'success': False, 'error': f'Task "{t.get("ticket","?")}": What I Did Today is required.'}), 400
            if not (t.get('hours') or '').strip() or (t.get('hours') or '').strip() == 'N/A':
                return jsonify({'success': False, 'error': f'Task "{t.get("ticket","?")}": Hours Spent is required.'}), 400

        llm_result = call_llm(form_data, email)
        work_date  = form_data.get('date', get_working_date())
        email_html = build_email_html(work_date, form_data, llm_result, email)

        T = Query()
        for task in form_data.get('tasks', []):
            name = task.get('ticket', '').strip()
            if name and not tickets_table.search((T.name == name) & (T.user == email)):
                tickets_table.insert({'name': name, 'user': email})

        teams_msgs = llm_result.get('teams_messages', [])
        # Fallback: if LLM returned old single string format
        if not teams_msgs:
            single = llm_result.get('teams_message', '')
            teams_msgs = [single] if single else []
        return jsonify({
            'success':        True,
            'teams_messages': teams_msgs,
            'teams_message':  '\n\n---\n\n'.join(teams_msgs),
            'email_html':     email_html,
            'email_subject':  llm_result.get('email_subject', f"ATTENDANCE: {work_date}")
        })
    except json.JSONDecodeError as e:
        return jsonify({'success': False, 'error': f'LLM returned invalid JSON. Try again. ({e})'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/save', methods=['POST'])
@login_required_api
def api_save():
    try:
        email        = get_current_user()
        data         = request.json
        form_data    = data.get('form_data', {})
        teams_msg    = data.get('teams_message', '')
        work_date    = form_data.get('date', get_working_date())
        date_key     = work_date.replace('/', '-')

        attachment, attach_err = _decode_attachment(data.get('attachment'))
        if attach_err:
            return jsonify({'success': False, 'error': attach_err}), 400

        drive_result = save_to_drive(email, date_key, work_date, form_data, teams_msg, attachment=attachment)

        R = Query()
        record = {
            'user':            email,
            'date':            date_key,
            'display_date':    work_date,
            'teams_drive_id':  drive_result['teams_drive_id'],
            'mail_drive_id':   drive_result['mail_drive_id'],
            'teams_drive_url': drive_result['teams_drive_url'],
            'mail_drive_url':  drive_result['mail_drive_url'],
            # Cleared to '' when no attachment this save, so a stale link from a
            # previous save for the same day never lingers once removed.
            'attachment_drive_id':  drive_result.get('attachment_drive_id', ''),
            'attachment_drive_url': drive_result.get('attachment_drive_url', ''),
            'attachment_filename':  drive_result.get('attachment_filename', ''),
            'task_count':      len(form_data.get('tasks', [])),
            'tickets':         [t.get('ticket','') for t in form_data.get('tasks', [])],
            'saved_at':        datetime.now().isoformat()
        }
        if records_table.search((R.date == date_key) & (R.user == email)):
            records_table.update(record, (R.date == date_key) & (R.user == email))
        else:
            records_table.insert(record)
        sync_db_to_drive()

        return jsonify({'success': True, 'message': 'Saved to Google Drive!', 'drive': drive_result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/send_email', methods=['POST'])
@login_required_api
def api_send_email():
    try:
        email = get_current_user()
        s     = get_user_settings(email)
        gmail_user = s.get('gmail_user', '')
        gmail_pass = s.get('gmail_app_password', '')
        to_email   = s.get('to_email', '')
        cc_raw     = s.get('cc_emails', '')
        cc_emails  = [x.strip() for x in cc_raw.split(',') if x.strip()] if isinstance(cc_raw, str) else cc_raw

        if not gmail_user or not gmail_pass:
            return jsonify({'success': False, 'error': 'Gmail not configured. Go to Settings.'}), 400
        if not to_email:
            return jsonify({'success': False, 'error': 'To email not configured. Go to Settings.'}), 400

        data          = request.json
        email_html    = data.get('email_html', '')
        email_subject = data.get('email_subject', 'Daily Work Update')

        attachment, attach_err = _decode_attachment(data.get('attachment'))
        if attach_err:
            return jsonify({'success': False, 'error': attach_err}), 400

        full_html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"></head>
<body style="margin:0;padding:20px;font-family:Arial,sans-serif;">{email_html}</body></html>"""

        sender_name = s.get('user_name') or s.get('sender_name') or s.get('employee_name') or ''
        if not sender_name and gmail_user:
            sender_name = gmail_user.split('@')[0].replace('.', ' ').replace('_', ' ').title()

        all_recipients = [to_email] + cc_emails
        try:
            send_via_gmail_api(email, all_recipients, email_subject, full_html, attachment=attachment, sender_name=sender_name)
            return jsonify({'success': True, 'message': 'Email sent! Check your Gmail Sent folder.'})
        except Exception as api_err:
            print(f"⚠️  Gmail API error for {email}: {api_err}")
            err = str(api_err)
            # Try SMTP fallback (works on local dev, fails on Railway — reported to user)
            try:
                _send_raw_email(gmail_user, gmail_pass, to_email,
                                email_subject, full_html, cc_emails=cc_emails, attachment=attachment, sender_name=sender_name)
                return jsonify({'success': True, 'message': 'Email sent!'})
            except Exception:
                return jsonify({'success': False, 'error': f'Email failed: {err}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/history')
@login_required_api
def api_history():
    email = get_current_user()
    R     = Query()
    records = records_table.search(R.user == email)

    def _date_sort_key(rec):
        d = rec.get('date', '')
        parts = d.replace('-', '/').split('/')
        if len(parts) == 3:
            return (parts[0] + parts[1] + parts[2]) if len(parts[0]) == 4 else (parts[2] + parts[1] + parts[0])
        return d

    records.sort(key=_date_sort_key, reverse=True)
    return jsonify({'records': records})

@app.route('/api/history/rebuild', methods=['POST'])
@login_required_api
def api_history_rebuild():
    email = get_current_user()
    try:
        rebuilt = rebuild_history_from_drive(email)
        return jsonify({'success': True, 'rebuilt': rebuilt,
                        'message': f'Restored {rebuilt} report(s) from Drive.'
                                   if rebuilt else 'All reports already indexed.'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/history/<date_key>')
@login_required_api
def api_history_detail(date_key):
    email = get_current_user()
    R     = Query()
    found = records_table.search((R.date == date_key) & (R.user == email))
    if not found:
        return jsonify({'error': 'Not found'}), 404
    record = found[0]
    teams_content = ''
    file_id = record.get('teams_drive_id', '')
    if file_id:
        try:
            teams_content = fetch_teams_from_drive(email, file_id)
        except Exception as e:
            teams_content = f'[Could not fetch from Drive: {e}]'
    return jsonify({'record': record, 'teams_content': teams_content})

@app.route('/api/download/<file_type>/<date_key>')
@login_required_api
def api_download(file_type, date_key):
    email = get_current_user()
    R     = Query()
    found = records_table.search((R.date == date_key) & (R.user == email))
    if not found:
        abort(404)
    record = found[0]
    if file_type == 'teams':
        url = record.get('teams_drive_url')
    elif file_type == 'attachment':
        url = record.get('attachment_drive_url')
    else:
        url = record.get('mail_drive_url')
    if not url:
        abort(404)
    return redirect(url)


@app.route('/admin/debug')
def admin_debug():
    if session.get('user_email') != ADMIN_EMAIL:
        return "Not authorized.", 403
    S = Query()
    admin_row  = settings_table.search(S.email == ADMIN_EMAIL)
    has_token  = bool(admin_row and admin_row[0].get('drive_token_json'))
    has_groq   = bool(admin_row and admin_row[0].get('groq_api_key'))
    has_gmail  = bool(admin_row and admin_row[0].get('gmail_user'))
    token_scopes = ''
    if has_token:
        try:
            t = json.loads(admin_row[0]['drive_token_json'])
            token_scopes = str(t.get('scopes', []))
            expiry       = t.get('expiry', 'unknown')
        except Exception:
            token_scopes = 'parse error'
            expiry = 'unknown'
    else:
        expiry = 'N/A'
    users = users_table.all()
    env_token = bool(os.getenv('ADMIN_DRIVE_TOKEN_JSON', '').strip())
    return f'''<pre style="padding:20px;font-family:monospace;font-size:13px">
=== VERCEL DB STATE ===
Admin settings row found : {bool(admin_row)}
  has drive token        : {has_token}
  token scopes           : {token_scopes}
  token expiry           : {expiry}
  has groq key           : {has_groq}
  has gmail              : {has_gmail}
ADMIN_DRIVE_TOKEN_JSON   : {"SET" if env_token else "MISSING"}
Users ({len(users)})     : {[u.get("email")+"="+u.get("status","?") for u in users]}
</pre>
<form method="post" action="/admin/force-sync">
  <button style="padding:10px 20px;background:#2563eb;color:#fff;border:none;border-radius:6px;cursor:pointer;font-size:14px">
    Force Sync from Drive now
  </button>
</form>'''


@app.route('/admin/force-sync', methods=['POST'])
def admin_force_sync():
    if session.get('user_email') != ADMIN_EMAIL:
        return "Not authorized.", 403
    sync_db_from_drive()
    S = Query()
    row = settings_table.search(S.email == ADMIN_EMAIL)
    has_settings = bool(row and row.get('groq_api_key') if row else False)
    return f'<pre style="padding:20px">Sync done.\nAdmin settings in DB: {bool(row)}\nHas groq: {bool(row and row[0].get("groq_api_key"))}\nHas drive token: {bool(row and row[0].get("drive_token_json"))}</pre>'


@app.route('/admin/remove-unknown')
def admin_remove_unknown():
    """Admin-only: remove all users with unknown/pending status and sync to Drive."""
    if session.get('user_email') != ADMIN_EMAIL:
        return "Not authorized.", 403
    U = Query()
    removed = users_table.remove((U.status == 'unknown') | (U.status == 'pending'))
    settings_table.remove((U.email == 'wp112.department@gmail.com') |
                          (U.email == 'barotb076@gmail.com'))
    sync_db_to_drive()
    remaining = [r.get('email') for r in users_table.all()]
    return f"Done. Removed unknown/pending users. Remaining: {remaining}"


@app.route('/admin/clear-users')
def admin_clear_users():
    """Admin-only: remove all non-admin users and sync to Drive."""
    if session.get('user_email') != ADMIN_EMAIL:
        return "Not authorized.", 403
    U = Query()
    users_table.remove(U.email != ADMIN_EMAIL)
    settings_table.remove(U.email != ADMIN_EMAIL)
    otps_table.truncate()
    sync_db_to_drive()
    return "Done. All non-admin users removed and synced to Drive."


@app.route('/admin/export-token')
def admin_export_token():
    """Admin-only: export Drive token JSON for ADMIN_DRIVE_TOKEN_JSON env var setup."""
    if session.get('user_email') != ADMIN_EMAIL:
        return "Not authorized. Log in as admin first.", 403
    S = Query()
    found = settings_table.search(S.email == ADMIN_EMAIL)
    if not found or not found[0].get('drive_token_json'):
        return "No token found. Connect Google Drive in Settings first.", 400
    token = found[0]['drive_token_json']
    return f'''<!doctype html><html><head><title>Export Token</title>
<style>body{{font-family:sans-serif;padding:30px;max-width:900px;margin:auto}}
textarea{{width:100%;height:180px;font-family:monospace;font-size:12px;padding:10px}}
.btn{{background:#2563eb;color:#fff;border:none;padding:10px 20px;border-radius:6px;cursor:pointer;font-size:14px}}
</style></head><body>
<h2>Admin Drive Token</h2>
<p>Copy the entire value below and set it as the <code>ADMIN_DRIVE_TOKEN_JSON</code>
environment variable on Railway or Vercel. After this, redeploys will automatically
restore all user credentials from Drive — no one needs to reconnect.</p>
<textarea id="t" readonly>{token}</textarea><br><br>
<button class="btn" onclick="navigator.clipboard.writeText(document.getElementById('t').value);this.textContent='Copied!'">Copy to Clipboard</button>
</body></html>'''


if __name__ == '__main__':
    print("\n" + "="*50)
    print("  Daily Work Update Generator")
    print("  http://localhost:5000")
    print("="*50 + "\n")
    debug = os.getenv('FLASK_DEBUG', 'false').lower() == 'true'
    port = int(os.getenv('PORT', 5000))
    app.run(debug=debug, host='0.0.0.0', port=port)
